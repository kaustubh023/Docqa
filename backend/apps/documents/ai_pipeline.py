import os
from functools import lru_cache

import pandas as pd
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from django.conf import settings
from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

PERSIST_DIRECTORY = os.path.join(settings.BASE_DIR, "vector_db")
EMBEDDING_CACHE_BASE_DIR = (
    os.getenv("DOCQA_EMBEDDINGS_CACHE_DIR")
    or os.path.join(PERSIST_DIRECTORY, ".hf_cache")
)
HF_HOME_DIR = os.path.abspath(EMBEDDING_CACHE_BASE_DIR)
TRANSFORMERS_CACHE_DIR = os.path.join(HF_HOME_DIR, "transformers")
SENTENCE_TRANSFORMERS_CACHE_DIR = os.path.join(HF_HOME_DIR, "sentence_transformers")

# Force Hugging Face related caches to use a writable in-project location.
os.environ["HF_HOME"] = HF_HOME_DIR
os.environ["TRANSFORMERS_CACHE"] = TRANSFORMERS_CACHE_DIR
os.environ["SENTENCE_TRANSFORMERS_HOME"] = SENTENCE_TRANSFORMERS_CACHE_DIR


@lru_cache(maxsize=1)
def get_embeddings():
    """Lazily initialize embeddings so app startup does not require model download."""
    os.makedirs(HF_HOME_DIR, exist_ok=True)
    os.makedirs(TRANSFORMERS_CACHE_DIR, exist_ok=True)
    os.makedirs(SENTENCE_TRANSFORMERS_CACHE_DIR, exist_ok=True)
    return HuggingFaceEmbeddings(
        model_name="all-MiniLM-L6-v2",
        cache_folder=SENTENCE_TRANSFORMERS_CACHE_DIR,
    )


def delete_document_embeddings(document_id):
    """Delete all vector chunks associated with a document id."""
    if not document_id:
        return
    try:
        vectorstore = Chroma(
            persist_directory=PERSIST_DIRECTORY,
            embedding_function=get_embeddings(),
        )
        vectorstore.delete(where={"document_id": str(document_id)})
    except Exception as error:
        print(f"VECTOR DELETE ERROR for document {document_id}: {error}")


def extract_text(file_path):
    """Extract text from supported file types."""
    ext = os.path.splitext(file_path)[1].lower()
    text = ""

    try:
        if ext == ".pdf":
            reader = PdfReader(file_path)
            for page in reader.pages:
                text += (page.extract_text() or "") + "\n"

        elif ext == ".docx":
            doc = DocxDocument(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"

        elif ext == ".txt":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()

        elif ext == ".csv":
            # Try common encodings so exported CSVs don't fail parsing.
            last_error = None
            for encoding in ("utf-8-sig", "utf-8", "latin-1"):
                try:
                    df = pd.read_csv(file_path, encoding=encoding, on_bad_lines="skip")
                    rows = []
                    for index, row in df.iterrows():
                        row_data = [
                            f"{col} is {val}" for col, val in row.items() if pd.notna(val)
                        ]
                        if row_data:
                            rows.append(f"Record {index + 1}: " + ", ".join(row_data))
                    text = "\n".join(rows)
                    break
                except Exception as csv_error:
                    last_error = csv_error
            if not text and last_error:
                print(f"CSV Extraction Error: {last_error}")
                return ""

    except Exception as error:
        print(f"General Extraction Error: {error}")
        return ""

    return text


def process_document(file_path, source_filename=None, document_id=None, user_id=None):
    """Chunk extracted text and index it in Chroma."""
    try:
        raw_text = extract_text(file_path)
        if not raw_text or not raw_text.strip():
            print(f"EXTRACTION FAILED: no text found in {file_path}")
            return False

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=100)
        chunks = text_splitter.split_text(raw_text)
        if not chunks:
            print(f"CHUNKING FAILED: no chunks generated for {file_path}")
            return False

        if document_id:
            # Prevent duplicate chunk copies for retry/reprocess of the same document.
            delete_document_embeddings(document_id)

        # Keep filename metadata for backward compatibility with existing code/history.
        source_name = source_filename or os.path.basename(file_path)
        metadata = {"source": source_name}
        if document_id:
            metadata["document_id"] = str(document_id)
        if user_id:
            metadata["user_id"] = str(user_id)

        docs = [Document(page_content=chunk, metadata=metadata) for chunk in chunks]

        Chroma.from_documents(
            documents=docs,
            embedding=get_embeddings(),
            persist_directory=PERSIST_DIRECTORY,
        )
        print(f"Indexed {len(chunks)} chunks for source: {source_name}")
        return True
    except Exception as error:
        print(f"PIPELINE CRITICAL ERROR: {error}")
        return False
